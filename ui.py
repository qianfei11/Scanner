# -*- coding=utf-8 -*-

from tkinter import *
import matplotlib

matplotlib.use('TkAgg')
from all_information import *
from scan_ip import *
from PIL import Image, ImageTk
from docx import Document
import time
import tkinter.messagebox as msg
from tkinter import filedialog
from tkinter import ttk

BANNER = '开发人员：XXX\t联系方式：XXXXXXXXXXX\t开发时间：XXXX.XX'


class MY_GUI:
    def __init__(self, parent_init_name):
        self.parent_init_name = parent_init_name
        self.ListIP = []
        self.ListInformation = []
        self.cve_list = []
        self.cve_CheckVar = IntVar()
        self.diy_CheckVar = IntVar()
        self.ftp1_CheckVar = IntVar()
        self.ftp2_CheckVar = IntVar()
        self.ftp3_CheckVar = IntVar()
        self.mysql1_CheckVar = IntVar()
        self.mysql2_CheckVar = IntVar()
        self.mysql3_CheckVar = IntVar()
        self.mysql4_CheckVar = IntVar()
        self.mongodb_CheckVar = IntVar()
        self.mongodb2_CheckVar = IntVar()
        self.have_cve_bug = 0
        self.have_ftp_bug = 0
        self.have_mysql_bug = 0
        self.have_mongodb_bug = 0
        self.scanned = 0
        self.filename = []

    def set_init_windows(self):
        # 定义开发窗口大小和形式
        self.parent_init_name.title('基于nmap的局域网网段扫描器')
        self.parent_init_name.geometry('700x900+10+10')
        self.parent_init_name.resizable(width=True, height=True)

        # self.parent_init_name['bg'] = '#000000'
        # self.parent_init_name.attributes('-alpha',0.1)

        # canvas = Canvas(self.parent_init_name)
        # im = Image.open('./233.jpg')
        # img = ImageTk.PhotoImage(im)
        # canvas.create_image(30,30,image = img)
        # canvas.pack()

        # ImageFrame = Frame(parent_init_name)
        # ImgFrame = LabelFrame(ImageFrame,text = '帮助')
        # im = Image.open('./3.jpg')
        # img = ImageTk.PhotoImage(im)
        # Label(ImgFrame, width=450,height=160,image=img).grid(row=0, column=0,columnspan=3)
        # ImgFrame.pack()
        # ImageFrame.pack()

        self.textframe = Frame(self.parent_init_name)
        self.developer = LabelFrame(self.textframe, text='开发信息')
        self.developer.pack(padx=10, pady=10)
        self.textframe.pack()
        # print(Image)
        self.im = Image.open('./banner.jpg')
        # 括号里为需要显示在图形化界面里的图片
        self.im = self.im.resize((560, 50))
        self.img = ImageTk.PhotoImage(self.im)
        Label(self.developer, width=560, height=50, image=self.img).grid(
            row=0, column=0, columnspan=1
        )
        Label(self.developer, width=70, height=1, bg='white', text=BANNER).grid(
            column=0, row=1
        )

        def clear_entry(event, entry):
            entry.delete(0, END)

        #  定义用户输入局域网内的IP范围和端口窗口，输入命令用来进行扫描
        self.user_input = Frame(self.parent_init_name)
        self.user_input_label = LabelFrame(self.user_input, text='扫描网段以及端口')
        self.user_input_label.pack(padx=10, pady=10)
        self.input_ip = StringVar()
        # self.input_ip.set('请输入需要扫描的IP')
        self.input_port = StringVar()
        # self.input_port.set('输入的端口')
        self.input_help = Label(
            self.user_input_label, width=70, height=1, bg='white', text='请输入扫描的网段以及端口'
        )
        # self.ip_label = Label(self.user_input_label, width=70, height=2)
        # self.L1 = Label(self.user_input_label,width=20, text='请输入需要扫描的IP：')
        self.entry_ip = Entry(
            self.user_input_label,
            width=70,
            bd=0,
            textvariable=self.input_ip,
            background='#A9A9A9',
        )
        self.entry_ip.insert(0, '请输入需要扫描的IP')
        self.entry_ip.bind(
            '<Button-1>', lambda event: clear_entry(event, self.entry_ip)
        )
        # self.port_label = Label(self.user_input_label, width=70, height=2)
        # self.L2 = Label(self.user_input_label,width=20, text='输入的端口：')
        self.entry_port = Entry(
            self.user_input_label,
            width=70,
            bd=0,
            textvariable=self.input_port,
            background='#A9A9A9',
        )
        self.entry_port.insert(0, '请输入需要扫描的端口')
        self.entry_port.bind(
            '<Button-1>', lambda event: clear_entry(event, self.entry_port)
        )
        # self.entry_ip.bind('<Return>', self.start_scan)
        # self.entry_port.bind('<Return>', self.start_scan)
        self.input_help.pack()
        # self.L1.pack(padx=20)
        self.entry_ip.pack()
        # self.L2.pack(padx=20)
        self.entry_port.pack()
        # self.ip_label.pack()
        # self.port_label.pack()
        # self.user_input_label.pack()
        self.user_input.pack()
        # start_scan(entry_ip.get(),entry_port.get())

        # 勾选框
        # self.cve_CheckVar = IntVar()
        # self.ftp_CheckVar = IntVar()
        # self.mysql_CheckVar = IntVar()
        # self.mongodb_CheckVar = IntVar()
        self.checkftpbox_input = Frame(self.parent_init_name)
        self.checkftpbox_input_label = LabelFrame(
            self.checkftpbox_input, text='勾选需要扫描的FTP相关内容'
        )
        self.checkftpbox_input_label.pack(padx=10, pady=10)
        self.ftp1_checkbox = Checkbutton(
            self.checkftpbox_input_label,
            text='FTP匿名登录漏洞',
            onvalue=1,
            variable=self.ftp1_CheckVar,
            offvalue=0,
            height=1,
            width=20,
        )
        self.ftp2_checkbox = Checkbutton(
            self.checkftpbox_input_label,
            text='FTP账户密码漏洞',
            onvalue=1,
            variable=self.ftp2_CheckVar,
            offvalue=0,
            height=1,
            width=20,
        )
        self.ftp3_checkbox = Checkbutton(
            self.checkftpbox_input_label,
            text='FTP服务信息',
            onvalue=1,
            variable=self.ftp3_CheckVar,
            offvalue=0,
            height=1,
            width=20,
        )
        self.ftp1_checkbox.pack(side='left')
        self.ftp2_checkbox.pack(side='left')
        self.ftp3_checkbox.pack(side='left')
        self.checkftpbox_input.pack()

        self.checkmysqlbox_input = Frame(self.parent_init_name)
        self.checkmysqlbox_input_label = LabelFrame(
            self.checkmysqlbox_input, text='勾选需要扫描的MYSQL相关内容'
        )
        self.checkmysqlbox_input_label.pack(padx=10, pady=10)
        self.mysql1_checkbox = Checkbutton(
            self.checkmysqlbox_input_label,
            text='MYSQL空口令漏洞',
            onvalue=1,
            variable=self.mysql1_CheckVar,
            offvalue=0,
            height=1,
            width=14,
        )
        self.mysql2_checkbox = Checkbutton(
            self.checkmysqlbox_input_label,
            text='MYSQL账户密码漏洞',
            onvalue=1,
            variable=self.mysql2_CheckVar,
            offvalue=0,
            height=1,
            width=16,
        )
        self.mysql3_checkbox = Checkbutton(
            self.checkmysqlbox_input_label,
            text='MYSQL用户信息',
            onvalue=1,
            variable=self.mysql3_CheckVar,
            offvalue=0,
            height=1,
            width=12,
        )
        self.mysql4_checkbox = Checkbutton(
            self.checkmysqlbox_input_label,
            text='MYSQL普通信息',
            onvalue=1,
            variable=self.mysql4_CheckVar,
            offvalue=0,
            height=1,
            width=14,
        )
        self.mysql1_checkbox.pack(side='left')
        self.mysql2_checkbox.pack(side='left')
        self.mysql3_checkbox.pack(side='left')
        self.mysql4_checkbox.pack(side='left')
        self.checkmysqlbox_input.pack()

        self.checkbox_input = Frame(self.parent_init_name)
        self.checkbox_input_label = LabelFrame(
            self.checkbox_input, text='勾选需要扫描的MONGODB相关内容'
        )
        self.checkbox_input_label.pack(padx=10, pady=10)
        self.mongodb_checkbox = Checkbutton(
            self.checkbox_input_label,
            text='MONGODB未授权访问漏洞',
            onvalue=1,
            variable=self.mongodb_CheckVar,
            offvalue=0,
            height=1,
            width=31,
        )
        self.mongodb2_checkbox = Checkbutton(
            self.checkbox_input_label,
            text='MONGODB账户密码漏洞',
            onvalue=1,
            variable=self.mongodb2_CheckVar,
            offvalue=0,
            height=1,
            width=32,
        )
        self.mongodb_checkbox.pack(side='left')
        self.mongodb2_checkbox.pack(side='left')
        self.checkbox_input.pack()

        self.checkbox_input = Frame(self.parent_init_name)
        self.checkbox_input_label = LabelFrame(
            self.checkbox_input, text='勾选需要扫描的CVE相关内容或者自定义脚本上传的相关内容'
        )
        self.checkbox_input_label.pack(padx=10, pady=10)
        self.cve_checkbox = Checkbutton(
            self.checkbox_input_label,
            text='主机所含的CVE漏洞相关信息',
            onvalue=1,
            variable=self.cve_CheckVar,
            offvalue=0,
            height=1,
            width=31,
        )
        self.diy_checkbox = Checkbutton(
            self.checkbox_input_label,
            text='上传自定义脚本所得相关信息',
            onvalue=1,
            variable=self.diy_CheckVar,
            offvalue=0,
            height=1,
            width=32,
        )
        self.cve_checkbox.pack(side='left')
        self.diy_checkbox.pack(side='left')
        self.checkbox_input.pack()

        # 下拉框
        # self.checkbox_input = Frame(self.parent_init_name)
        # self.checkbox_input_label = LabelFrame(self.checkbox_input, text='勾选需要扫描的信息')
        # self.checkbox_input_label.pack(padx=10,pady=10)
        # self.choice = ttk.Combobox(self.parent_init_name,textvariable=xVariabl)
        # self.choice.pack()

        # 查看扫描出来的局域网内主机ip详细信息
        self.IpFrame = Frame(self.parent_init_name)
        self.iplistframe = LabelFrame(self.IpFrame, text='局域网在线主机ip')
        self.iplistframe.pack(padx=10, pady=10)
        self.ip_sb = Scrollbar(self.iplistframe)
        self.ip_sb.pack(side=RIGHT, fill=Y)
        self.ip_sb2 = Scrollbar(self.iplistframe, orient=HORIZONTAL)
        self.ip_sb2.pack(side=BOTTOM, fill=X, expand=True)
        self.listip = Listbox(
            self.iplistframe,
            height=2,
            width=68,
            bd=0,
            yscrollcommand=self.ip_sb.set,
            xscrollcommand=self.ip_sb2.set,
        )
        self.listip.bind('<Double-Button-1>', self.get_ip_information)
        # self.listip.bind('<Double-Button-2>', self.get_cve_information)
        for item in self.ListIP:
            self.listip.insert(END, item)
        self.listip.pack()
        self.IpFrame.pack()
        self.ip_sb.config(command=self.listip.yview)
        self.ip_sb2.config(command=self.listip.xview)

        # 主机的详细信息
        # 开放端口，os信息，防火墙，杀毒软件
        # 存在的安全漏洞
        self.Information = Frame(self.parent_init_name)
        self.ipinformation = LabelFrame(self.Information, text='局域网主机详细信息')
        self.ipinformation.pack(padx=10, pady=10)
        self.info_sb = Scrollbar(self.ipinformation)
        self.info_sb.pack(side=RIGHT, fill=Y)
        self.info_sb2 = Scrollbar(self.ipinformation, orient=HORIZONTAL)
        self.info_sb2.pack(side=BOTTOM, fill=X, expand=True)
        self.listinformation = Listbox(
            self.ipinformation,
            height=2,
            width=68,
            bd=0,
            yscrollcommand=self.info_sb.set,
            xscrollcommand=self.info_sb2.set,
        )
        for item in self.ListInformation:
            # if '存在' in item:
            #     self.listinformation = Listbox(self.ipinformation,fg = 'red')
            self.listinformation.insert(END, item)
        self.listinformation.pack()
        self.Information.pack()
        self.info_sb.config(command=self.listinformation.yview)
        self.info_sb2.config(command=self.listinformation.xview)

        # cve列表界面
        self.cve_information = Frame(self.parent_init_name)
        self.cveinformation = LabelFrame(self.cve_information, text='cve信息列表')
        self.cveinformation.pack(padx=10, pady=10)
        self.cve_sb = Scrollbar(self.cveinformation)
        self.cve_sb.pack(side=RIGHT, fill=Y)
        self.cve_sb2 = Scrollbar(self.cveinformation, orient=HORIZONTAL)
        self.cve_sb2.pack(side=BOTTOM, fill=X, expand=True)
        self.listCveinformation = Listbox(
            self.cveinformation,
            height=2,
            width=68,
            bd=0,
            yscrollcommand=self.cve_sb.set,
            xscrollcommand=self.cve_sb2.set,
        )
        for item in self.cve_list:
            self.listCveinformation.insert(END, item)
        self.listCveinformation.pack()
        self.cve_information.pack()
        self.cve_sb.config(command=self.listCveinformation.yview)
        self.cve_sb2.config(command=self.listCveinformation.xview)

        # UI界面按钮
        self.buttonframe = Frame(self.parent_init_name)
        Button(self.buttonframe, text='开始', command=self.start_scan).grid(
            column=0, row=0
        )
        Button(self.buttonframe, text='退出', command=self.exitui).grid(column=4, row=0)
        Button(self.buttonframe, text='帮助', command=self.help).grid(column=5, row=0)
        Button(self.buttonframe, text='建议', command=self.advice).grid(column=1, row=0)
        Button(self.buttonframe, text='导出结果', command=self.report).grid(column=3, row=0)
        Button(self.buttonframe, text='上传自定义文件', command=self.upload_file).grid(
            column=2, row=0
        )
        self.buttonframe.pack()

    def start_scan(self):
        # self.start_process_bar()
        # self.bar.start()

        # for item in self.listip:
        # self.listip.delete(0,END)
        # print(self.listip)
        # print(self.Listip)

        del ip_alive_list[:]
        del all_ip_all_information[:]

        ip_input = self.entry_ip.get()
        port_input = self.entry_port.get()
        # 对用户输入的端口信息进行处理
        listip = getip(ip_input)
        startport, endport = getport(port_input)
        # 使用多线程进程扫描在线的ip
        threads_find_ip(listip)

        if len(ip_alive_list) == 0:
            print('没有扫描到存在的IP地址')
        else:
            # 对在线的ip列表进行排序
            sort_alive_ip(ip_alive_list)
            # 多线程对每一个ip的详细进行扫描
            a = all_infromation(
                startport, endport, ip_alive_list, self.get_vars(), self.filename
            )
            a.Threads()

            (
                self.have_cve_bug,
                self.have_ftp_bug,
                self.have_mysql_bug,
                self.have_mongodb_bug,
            ) = a.get_bug_status()

            self.listip.delete(0, END)
            for i in ip_alive_list:
                self.listip.insert(END, i)

            self.scanned = 1

        # self.stop_process_bar()

    # 获取扫描IP的详细信息
    def get_ip_information(self, event):
        index_ip = self.listip.curselection()
        # ip_for_list = ip_alive_list[index_ip[0]]
        self.all_information = all_ip_all_information[index_ip[0]]
        self.listinformation.delete(0, END)
        for item in self.all_information:
            self.listinformation.insert(END, item)
            if '存在' in item:
                self.listinformation.itemconfig(END, fg='red')
            # else:
            #     self.listinformation.itemconfig(END,fg='red')

        if all_ip_cve_information != []:
            self.all_cveinformation = all_ip_cve_information[index_ip[0]]
            self.listCveinformation.delete(0, END)
            for item in self.all_cveinformation:
                self.listCveinformation.insert(END, item)
                if 'CVE' in item:
                    self.listCveinformation.itemconfig(END, fg='red')
        else:
            self.all_cveinformation = []

    def get_cve_information(self, event):
        index_ip = self.listip.curselection()
        # ip_for_list = ip_alive_list[index_ip[0]]
        self.all_cveinformation = all_ip_cve_information[index_ip[0]]
        self.listCveinformation.delete(0, END)
        for item in self.all_cveinformation:
            self.listCveinformation.insert(END, item)

    # 帮助中的注释框
    def help(self):
        top = Toplevel()
        top.title('帮助')
        top.geometry('600x600')

        ImageFrame = Frame(top)
        ImgFrame = LabelFrame(ImageFrame, text='帮助')
        # print(Image)
        im = Image.open('./nmap.png')
        # 括号里为需要显示在图形化界面里的图片
        im = im.resize((450, 160))
        img = ImageTk.PhotoImage(im)
        Label(ImgFrame, width=450, height=160, image=img).grid(
            row=0, column=0, columnspan=3
        )
        ImgFrame.pack()
        ImageFrame.pack()

        textframe = Frame(top)
        developer = LabelFrame(textframe, text='使用教程')
        developer.pack(padx=10, pady=10)
        textframe.pack()
        Label(
            developer, width=50, height=2, bg='white', text='环境配置：python，python-nmap\n'
        ).grid(column=0, row=0)
        Label(
            developer,
            width=50,
            height=2,
            bg='white',
            text='开始按钮：点击即可开始扫描主机所在的网段的所有在线的主机\n',
        ).grid(column=0, row=1)
        Label(
            developer,
            width=50,
            height=2,
            bg='white',
            text='退出按钮：点击即可退出扫描程序                      \n',
        ).grid(column=0, row=2)
        Label(
            developer,
            width=50,
            height=2,
            bg='white',
            text='帮助按钮：提示程序如何进行使用                      \n',
        ).grid(column=0, row=3)
        Label(
            developer,
            width=50,
            height=2,
            bg='white',
            text='局域网主机ip：显示该网段的在线主机点击可查看特定主机的详细信息\n',
        ).grid(column=0, row=4)

        top.mainloop()

    def exitui(self):
        exit(0)

    def get_vars(self):
        return [
            self.cve_CheckVar.get(),
            self.ftp1_CheckVar.get(),
            self.ftp2_CheckVar.get(),
            self.ftp3_CheckVar.get(),
            self.mysql1_CheckVar.get(),
            self.mysql2_CheckVar.get(),
            self.mysql3_CheckVar.get(),
            self.mysql4_CheckVar.get(),
            self.mongodb_CheckVar.get(),
            self.mongodb2_CheckVar.get(),
            self.diy_CheckVar.get(),
        ]

    def advice(self):
        top = Toplevel()
        top.title('安全加固建议')
        top.geometry('600x600')
        textframe = Frame(top)
        print(self.scanned)
        if self.scanned == 1:
            if self.ftp1_CheckVar.get() == 1 and self.have_ftp_bug == 1:
                developer = LabelFrame(textframe, text='FTP漏洞加固建议')
                developer.pack(padx=10, pady=10)
                textframe.pack()
                text = Text(developer, width=100, height=8)
                text.pack()
                text.insert(INSERT, 'FTP漏洞加固：\n')
                text.insert(
                    END,
                    '\t将配置文件中找到anonymous_enable，将该参数配置为 NO 表示禁止匿名登录，必须要创建用户认证后才能登录 FTP 服务。\n',
                )
            if self.ftp2_CheckVar.get() == 1 and self.have_ftp_bug == 1:
                developer = LabelFrame(textframe, text='FTP漏洞加固建议')
                developer.pack(padx=10, pady=10)
                textframe.pack()
                text = Text(developer, width=100, height=8)
                text.pack()
                text.insert(INSERT, 'FTP漏洞加固：\n')
                text.insert(
                    END, '\t自我检查FTP中的账户密码信息，创建管理员账号进行管理，避免账户密码是简单的账户密码，甚至是空密码\n'
                )
            if self.mysql1_CheckVar.get() == 1 and self.have_mysql_bug == 1:
                developer2 = LabelFrame(textframe, text='MYSQL漏洞加固建议')
                developer2.pack(padx=10, pady=10)
                textframe.pack()
                text2 = Text(developer2, width=100, height=8)
                text2.pack()
                text2.insert(INSERT, 'MYSQL漏洞加固：\n')
                text2.insert(
                    END,
                    '\t缺省安装的MySQL的root用户是空密码的，为了安全起见，必须修改为强密码，所谓的强密码，至少8位，由字母、数字和符号组成的不规律密码。使用MySQL自带的命令mysaladmin修改root密码，同时也可以登陆数据库，修改数据库mysql下的user表的字段内容，修改方法如下所示：\n注意：安装的mapn默认的mysql密码是root\na.登录mysql /usr/local/mysql/bin/mysql -u root -p 123456\nb.进入到mysql控制台后你会看到四个数据库information_schema,test,mysql,ftp(这些是针对mapn而言的)修改数据库mysql中的user表：update user set password=password(abcdef) where user=root;\nc.#mysql> flush privileges; //强制刷新内存授权表，否则用的还是在内存缓冲的口令\n',
                )
            if self.mysql2_CheckVar.get() == 1 and self.have_mysql_bug == 1:
                developer2 = LabelFrame(textframe, text='MYSQL漏洞加固建议')
                developer2.pack(padx=10, pady=10)
                textframe.pack()
                text2 = Text(developer2, width=100, height=8)
                text2.pack()
                text2.insert(INSERT, 'MYSQL漏洞加固：\n')
                text2.insert(
                    END, '\t自我检查MYSQL数据库中的账户密码信息，创建管理员账号进行管理，避免账户密码是简单的账户密码，甚至是空密码\n'
                )
            if self.mongodb_CheckVar.get() == 1 and self.have_mongodb_bug == 1:
                developer3 = LabelFrame(textframe, text='MONGODB漏洞加固建议')
                developer3.pack(padx=10, pady=10)
                textframe.pack()
                text3 = Text(developer3, width=100, height=8)
                text3.pack()
                text3.insert(INSERT, 'MONGODB漏洞加固：\n')
                text3.insert(
                    END,
                    '\t启动基于角色的登录认证功能:MongoDB 3.0及以上版本启动时添加--auth参数开启认证访问，此时若数据库中无账号，本地登录则无权限进行任何操作，因此需要先以无认证的方式启动服务并创建系统用户管理员账号。将配置文件中找到auth，并将auth=true\n',
                )
            if self.mongodb2_CheckVar.get() == 1 and self.have_mongodb_bug == 1:
                developer3 = LabelFrame(textframe, text='MONGODB漏洞加固建议')
                developer3.pack(padx=10, pady=10)
                textframe.pack()
                text3 = Text(developer3, width=100, height=8)
                text3.pack()
                text3.insert(INSERT, 'MONGODB漏洞加固：\n')
                text3.insert(
                    END, '\t自我检查MONGODB数据库中的账户密码信息，创建管理员账号进行管理，避免账户密码是简单的账户密码，甚至是空密码\n'
                )
            # if self.mongodb_CheckVar == 0 and self.mysql_CheckVar == 0 and self.mongodb_CheckVar == 0:
            #     developer4 = LabelFrame(textframe, text='漏洞加固建议')
            #     developer4.pack(padx=10, pady=10)
            #     textframe.pack()
            #     text4 = Text(developer4,width=100,height = 8)
            #     text4.pack()
            #     text4.insert(INSERT,'漏洞加固建议：\n')
            #     text4.insert(END,' 您的主机很健康，不需要安全建议\n')
            # Label(developer, width=50, height=2, bg='white', text='CVE漏洞加固：\n').grid(column=0, row=0)
            # Label(developer, width=50, height=2, bg='white', text='FTP漏洞加固：\n').grid(column=0,row=5)
            # Label(developer, width=50, height=2, bg='white', wraplength = 40,text='将配置文件中找到anonymous_enable，将该参数配置为 NO 表示禁止匿名登录，必须要创建用户认证后才能登录 FTP 服务。\n').grid(column=0,row=6)
            # Label(developer, width=50, height=2, bg='white', text='MYSQL漏洞加固：\n').grid(column=0,row=10)
            # Label(developer, width=50, height=2, bg='white', text='MONGODB漏洞加固：\n').grid(column=0, row=15)
        else:
            developer4 = LabelFrame(textframe, text='建议')
            developer4.pack(padx=10, pady=10)
            textframe.pack()
            text4 = Text(developer4, width=100, height=8)
            text4.pack()
            text4.insert(INSERT, '建议：\n')
            text4.insert(END, '\t请先扫描您的主机！\n')

        top.mainloop()

    def report(self):
        top = Toplevel()
        top.title('导出文件')
        top.geometry('300x150')
        textframe = Frame(top)
        developer = LabelFrame(textframe, text='保存文件命名')
        developer.pack(padx=10, pady=10)
        textframe.pack()
        input_name = StringVar()
        input_title = Label(
            developer, width=70, height=2, bg='white', text='请输入需要保存的名字'
        )
        entry_name = Entry(
            developer, width=70, bd=0, textvariable=input_name, background='#A9A9A9'
        )
        input_title.pack()
        entry_name.pack()

        def sure():
            print(self.scanned)
            # if self.scanned == 0:
            if self.scanned == 1:
                today = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(
                    y='年', m='月', d='日'
                )
                print(today)
                docxname = entry_name.get()
                print('docxname:', docxname)
                # print(self.listinformation)
                document = Document()
                paragraph = document.add_paragraph(u'主机扫描文件报告：')
                # paragraph = document.add_paragraph(self.listinformation)
                for i in self.all_information:
                    i = i.decode('utf-8')
                    paragraph = document.add_paragraph(i)

                for j in self.all_cveinformation:
                    j = j.decode('utf-8')
                    paragraph = document.add_paragraph(j)
                document.save(docxname + '.docx')
                print('保存成功！')
                msg.showerror('保存成功', '保存成功！！！')
                # document.save( '123.docx')
            else:
                msg.showerror('输入错误', '请先进行扫描！！！')

        buttonframe = Frame(top)
        Button(buttonframe, text='确认', command=sure).grid(column=0, row=0)
        buttonframe.pack()

        # print(self.scanned)
        # if self.scanned == 1:
        #     if self.ftp1_CheckVar.get() == 1 and self.have_ftp_bug == 1:
        #         developer = LabelFrame(textframe, text='FTP漏洞加固建议')
        #         developer.pack(padx=10, pady=10)
        #         textframe.pack()
        #         text = Text(developer,width=100,height = 8)
        #         text.pack()
        #         text.insert(INSERT,'FTP漏洞加固：\n')
        #         text.insert(END,'   将配置文件中找到anonymous_enable，将该参数配置为 NO 表示禁止匿名登录，必须要创建用户认证后才能登录 FTP 服务。\n')

        #     if self.mysql1_CheckVar.get() == 1 and self.have_mysql_bug == 1:
        #         developer2 = LabelFrame(textframe, text='MYSQL漏洞加固建议')
        #         developer2.pack(padx=10, pady=10)
        #         textframe.pack()
        #         text2 = Text(developer2,width=100,height = 8)
        #         text2.pack()
        #         text2.insert(INSERT,'MYSQL漏洞加固：\n')
        #         text2.insert(END,'  缺省安装的MySQL的root用户是空密码的，为了安全起见，必须修改为强密码，所谓的强密码，至少8位，由字母、数字和符号组成的不规律密码。使用MySQL自带的命令mysaladmin修改root密码，同时也可以登陆数据库，修改数据库mysql下的user表的字段内容，修改方法如下所示：\n注意：安装的mapn默认的mysql密码是root\na.登录mysql /usr/local/mysql/bin/mysql -u root -p 123456\nb.进入到mysql控制台后你会看到四个数据库information_schema,test,mysql,ftp(这些是针对mapn而言的)修改数据库mysql中的user表：update user set password=password(abcdef) where user=root;\nc.#mysql> flush privileges; //强制刷新内存授权表，否则用的还是在内存缓冲的口令\n')

        #     if self.mongodb_CheckVar.get() == 1 and self.have_mongodb_bug == 1:
        #         developer3 = LabelFrame(textframe, text='MONGODB漏洞加固建议')
        #         developer3.pack(padx=10, pady=10)
        #         textframe.pack()
        #         text3 = Text(developer3,width=100,height = 8)
        #         text3.pack()
        #         text3.insert(INSERT,'MONGODB漏洞加固：\n')
        #         text3.insert(END,'  启动基于角色的登录认证功能:MongoDB 3.0及以上版本启动时添加--auth参数开启认证访问，此时若数据库中无账号，本地登录则无权限进行任何操作，因此需要先以无认证的方式启动服务并创建系统用户管理员账号。将配置文件中找到auth，并将auth=true\n')
        # if self.mongodb_CheckVar == 0 and self.mysql_CheckVar == 0 and self.mongodb_CheckVar == 0:
        #     developer4 = LabelFrame(textframe, text='漏洞加固建议')
        #     developer4.pack(padx=10, pady=10)
        #     textframe.pack()
        #     text4 = Text(developer4,width=100,height = 8)
        #     text4.pack()
        #     text4.insert(INSERT,'漏洞加固建议：\n')
        #     text4.insert(END,' 您的主机很健康，不需要安全建议\n')

        # Label(developer, width=50, height=2, bg='white', text='CVE漏洞加固：\n').grid(column=0, row=0)
        # Label(developer, width=50, height=2, bg='white', text='FTP漏洞加固：\n').grid(column=0,row=5)
        # Label(developer, width=50, height=2, bg='white', wraplength = 40,text='将配置文件中找到anonymous_enable，将该参数配置为 NO 表示禁止匿名登录，必须要创建用户认证后才能登录 FTP 服务。\n').grid(column=0,row=6)

        # Label(developer, width=50, height=2, bg='white', text='MYSQL漏洞加固：\n').grid(column=0,row=10)
        # Label(developer, width=50, height=2, bg='white', text='MONGODB漏洞加固：\n').grid(column=0, row=15)

        # else:
        #     developer4 = LabelFrame(textframe, text='漏洞加固建议')
        #     developer4.pack(padx=10, pady=10)
        #     textframe.pack()
        #     text4 = Text(developer4,width=100,height = 8)
        #     text4.pack()
        #     text4.insert(INSERT,'漏洞加固建议：\n')
        #     text4.insert(END,'  请先扫描您的主机！\n')

        top.mainloop()

    def upload_file(self):
        selectFile = (
            filedialog.askopenfilename()
        )  # askopenfilename 1次上传1个；askopenfilenames1次上传多个
        # entry1.insert(0, selectFile)
        self.filename = selectFile.split('/')[-1]
        print(self.filename)
        if '.nse' in self.filename:
            with open(selectFile, 'rb') as f:
                data = f.read()
                # print(data)
            with open('/usr/share/nmap/scripts/' + self.filename, 'wb') as f:
                f.write(data)
            msg.showerror('上传成功', '上传成功！！！')
        else:
            msg.showerror('上传错误', '请上传.nse结尾的脚本文件！！！')

        # return self.filename

    def start_process_bar(self):
        self.process = Tk()
        self.process.geometry('150x120')
        self.bar = ttk.Progressbar(
            self.process, length=200, mode='indeterminate', orient=HORIZONTAL
        )
        self.bar.pack(padx=5, pady=10)
        self.process.mainloop()

    def stop_process_bar(self):
        self.bar.stop()
        self.process.quit()


def start():
    init_windows = Tk()
    ZMJ_PORTAL = MY_GUI(init_windows)
    ZMJ_PORTAL.set_init_windows()
    init_windows.mainloop()
